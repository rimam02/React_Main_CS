from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT_PATH = "JWT Security Laboratory Report.docx"
LOGO_PATH = "/Users/rimamaji/Desktop/ITM Logo.png"
SCREENSHOT_DIR = "/private/tmp/jwt_report_assets"


def set_cell_text(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return run


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_centered_paragraph(doc, text="", size=12, bold=False, after=8):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=after)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=12 if level == 1 else 8, after=8)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=8)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(paragraph, after=4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(paragraph, after=4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=2, after=10)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)
    return paragraph


def set_document_layout(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    for style_name in ["List Bullet", "List Number"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(11)


def add_page_number_footer(doc):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    run.font.name = "Arial"
    run.font.size = Pt(9)


def add_cover_page(doc):
    for _ in range(3):
        doc.add_paragraph()

    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(LOGO_PATH, width=Inches(3.9))

    for _ in range(5):
        doc.add_paragraph()

    add_centered_paragraph(doc, "School of Future Tech", size=22, bold=True, after=70)
    add_centered_paragraph(doc, "Case Study Report", size=22, bold=True, after=28)
    add_centered_paragraph(doc, "on", size=15, after=38)
    add_centered_paragraph(doc, "JWT (JSON Web Token) Security Laboratory", size=17, bold=True, after=22)
    add_centered_paragraph(doc, "by", size=15, after=34)
    add_centered_paragraph(doc, "Rima Maji", size=13, bold=True, after=16)
    add_centered_paragraph(doc, "Your Full Roll Number", size=13, after=0)
    doc.add_page_break()


def add_index_page(doc):
    add_centered_paragraph(doc, "Index", size=16, bold=True, after=18)
    items = [
        "Introduction to the Case Study.",
        "Problem Statement / Case Background (Abstract).",
        "Problem Statement / Case Study Design.",
        "Methods & Algorithms Technology Applied in the Problem Statement / Case Study.",
        "Problem Statement / Case Study Implementation Details and Snapshots.",
        "Problem Statement / Case Study Results and Conclusion.",
        "References.",
    ]
    for index, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph, after=7)
        run = paragraph.add_run(f"{index}.  {item}")
        run.font.name = "Arial"
        run.font.size = Pt(11)
    doc.add_page_break()


def add_report_content(doc):
    add_heading(doc, "1. Introduction to the Case Study")
    add_body(
        doc,
        "JSON Web Token (JWT) is widely used in modern web applications for sharing user identity and authorization information between client and server. A JWT is compact, easy to transfer, and contains encoded information in three parts: header, payload, and signature.",
    )
    add_body(
        doc,
        "This case study focuses on designing and implementing a JWT Security Laboratory using ReactJS. The laboratory helps beginners understand how JWTs are generated, decoded, verified, stored in browser localStorage, and tested against common security mistakes.",
    )
    add_body(
        doc,
        "The project is frontend-only and is built for educational purposes. It does not use any backend, database, Firebase, Redux, TypeScript, or Tailwind CSS. The main goal is to understand JWT structure and security concepts through interactive demonstrations.",
    )

    add_heading(doc, "2. Problem Statement / Case Background (Abstract)")
    add_heading(doc, "Background", level=2)
    add_body(
        doc,
        "Many students learn authentication theoretically but do not clearly understand what a token contains or why signature verification is important. JWT payloads are only encoded, not encrypted, so they can be decoded by anyone who has the token.",
    )
    add_body(
        doc,
        "A common security problem occurs when applications trust a token without verifying its signature. Attackers may try to change the payload, change the role value, or misuse weak JWT validation logic.",
    )
    add_heading(doc, "Abstract", level=2)
    add_body(
        doc,
        "This case study presents a ReactJS-based JWT Security Laboratory. The system allows the user to enter a JSON payload and secret key, generate a signed JWT token, decode the token header and payload, verify its signature, simulate none algorithm and payload tampering attacks, and store generated tokens in localStorage.",
    )
    add_body(
        doc,
        "The application uses the jose library for signing and verifying tokens, jwt-decode for decoding, React hooks for state management, and CSS for responsive light/dark user interface design.",
    )

    add_heading(doc, "3. Problem Statement / Case Study Design")
    add_body(
        doc,
        "The JWT Security Laboratory is designed as a modular React application. Each major feature is placed in a separate component to keep the code easy to understand and viva-friendly.",
    )
    add_numbered(doc, "JWT Token Generator")
    add_bullet(doc, "Accepts JSON payload and secret key from the user.")
    add_bullet(doc, "Generates a signed JWT using HS256 algorithm.")
    add_bullet(doc, "Displays success or error messages.")
    add_numbered(doc, "JWT Decoder")
    add_bullet(doc, "Accepts a JWT token pasted by the user.")
    add_bullet(doc, "Decodes and displays header and payload in readable JSON format.")
    add_numbered(doc, "JWT Signature Verification")
    add_bullet(doc, "Accepts token and secret key.")
    add_bullet(doc, "Shows whether the signature is valid or invalid.")
    add_numbered(doc, "JWT Attack Simulator")
    add_bullet(doc, "Demonstrates none algorithm attack by changing alg to none and removing the signature.")
    add_bullet(doc, "Demonstrates payload tampering by modifying the payload while keeping the old signature.")
    add_numbered(doc, "Token History and Information Panel")
    add_bullet(doc, "Stores generated tokens in localStorage with date and time.")
    add_bullet(doc, "Provides beginner-friendly theory about JWT risks and best practices.")

    add_heading(doc, "4. Methods & Algorithms Technology Applied in the Problem Statement / Case Study")
    add_heading(doc, "Key methods and algorithms:", level=2)
    add_numbered(doc, "JWT signing")
    add_bullet(doc, "The jose library signs the payload using the HS256 algorithm.")
    add_bullet(doc, "The secret key is converted into bytes using TextEncoder before signing.")
    add_numbered(doc, "JWT decoding")
    add_bullet(doc, "The jwt-decode library decodes the header and payload.")
    add_bullet(doc, "Decoding only reads the token and does not verify security.")
    add_numbered(doc, "Signature verification")
    add_bullet(doc, "jwtVerify from jose checks whether the token signature matches the secret key.")
    add_bullet(doc, "If the token or secret key is wrong, verification fails.")
    add_numbered(doc, "Attack demonstration methods")
    add_bullet(doc, "The none algorithm attack changes the token header to alg none and removes the signature.")
    add_bullet(doc, "Payload tampering changes payload data but keeps the original signature, so verification fails.")
    add_heading(doc, "Technology Stack Used for the Case", level=2)
    add_bullet(doc, "Frontend framework: ReactJS with Vite.")
    add_bullet(doc, "Programming language: JavaScript.")
    add_bullet(doc, "Styling: CSS only.")
    add_bullet(doc, "JWT libraries: jose and jwt-decode.")
    add_bullet(doc, "Storage: Browser localStorage.")
    add_bullet(doc, "Environment: VS Code and browser-based local development server.")

    add_heading(doc, "5. Problem Statement / Case Study Implementation Details and Snapshots.")
    add_body(
        doc,
        "The implementation is organized into reusable React functional components. The main files are inside the src folder. App.jsx controls the overall layout, theme state, and token history. The components folder contains separate files for generator, decoder, verifier, attack simulator, history, information panel, navbar, and theme toggle.",
    )
    add_numbered(doc, "App.jsx")
    add_bullet(doc, "Maintains light/dark theme using useState.")
    add_bullet(doc, "Stores token history using useState and localStorage.")
    add_bullet(doc, "Passes functions and data to child components through props.")
    add_numbered(doc, "TokenGenerator.jsx")
    add_bullet(doc, "Takes payload JSON and secret key from the user.")
    add_bullet(doc, "Uses SignJWT from jose to create the token.")
    add_bullet(doc, "Saves generated token into history.")
    add_numbered(doc, "TokenDecoder.jsx")
    add_bullet(doc, "Uses jwtDecode to decode header and payload.")
    add_bullet(doc, "Handles invalid token errors.")
    add_numbered(doc, "TokenVerifier.jsx")
    add_bullet(doc, "Uses jwtVerify to check the token signature.")
    add_bullet(doc, "Shows Valid Signature or Invalid Signature.")
    add_numbered(doc, "AttackSimulator.jsx")
    add_bullet(doc, "Creates modified tokens for educational attack demonstrations.")
    add_bullet(doc, "Explains why accepting unsafe tokens is dangerous.")
    add_numbered(doc, "TokenHistory.jsx and InfoPanel.jsx")
    add_bullet(doc, "TokenHistory displays localStorage saved tokens with date and time.")
    add_bullet(doc, "InfoPanel explains JWT concepts and best practices.")

    screenshots = [
        ("01-home.png", "Snapshot 1: Home screen of JWT Security Laboratory"),
        ("02-generated-token.png", "Snapshot 2: JWT token generated successfully"),
        ("03-decoder-verifier.png", "Snapshot 3: Decoded token and valid signature verification"),
        ("04-attack-simulator.png", "Snapshot 4: None algorithm attack simulator output"),
    ]
    for filename, caption in screenshots:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(f"{SCREENSHOT_DIR}/{filename}", width=Inches(5.8))
        add_caption(doc, caption)

    add_heading(doc, "6. Problem Statement / Case Study Results and Conclusion.")
    add_heading(doc, "Findings", level=2)
    add_bullet(doc, "The project successfully generates JWT tokens from user-entered JSON payload and secret key.")
    add_bullet(doc, "The decoder clearly shows that JWT header and payload can be read by anyone.")
    add_bullet(doc, "The verifier proves that a token should be trusted only after signature verification.")
    add_bullet(doc, "The attack simulator demonstrates that changing the payload breaks the signature.")
    add_bullet(doc, "localStorage successfully stores generated token history and theme preference.")
    add_heading(doc, "Conclusion", level=2)
    add_body(
        doc,
        "This case study demonstrates a complete beginner-level JWT Security Laboratory using ReactJS. The project explains JWT generation, decoding, verification, localStorage usage, and common JWT vulnerabilities through an interactive frontend interface.",
    )
    add_body(
        doc,
        "The project is suitable for first-year engineering learning because it uses simple React functional components, useState, useEffect, reusable components, CSS, and beginner-friendly code comments. In real-world applications, JWT signing and verification should be done on a secure backend, but this frontend-only project is useful for understanding the concepts clearly.",
    )

    add_heading(doc, "7. References")
    add_bullet(doc, "Official React documentation for functional components and hooks.")
    add_bullet(doc, "Official Vite documentation for React project setup and development server.")
    add_bullet(doc, "jose library documentation for JWT signing and verification.")
    add_bullet(doc, "jwt-decode library documentation for decoding JSON Web Tokens.")
    add_bullet(doc, "MDN Web Docs for localStorage, JavaScript, and browser APIs.")
    add_bullet(doc, "Educational articles and tutorials on JWT authentication, authorization, and token security.")


def main():
    doc = Document()
    set_document_layout(doc)
    add_cover_page(doc)
    add_index_page(doc)
    add_report_content(doc)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
